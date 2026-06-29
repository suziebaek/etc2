import streamlit as st
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Twips
import io
import re
import os

st.set_page_config(
    page_title="정기평가 양식 상호 변환 시스템",
    page_icon="📝",
    layout="centered"
)

# ==========================================
# [공통 유틸리티 함수]
# ==========================================
def is_originally_red(run):
    try:
        if run.font.color and run.font.color.rgb:
            hex_color = str(run.font.color.rgb).upper()
            if hex_color == "FF0000" or "255, 0, 0" in hex_color:
                return True
    except:
        pass
    return False

def should_skip_paragraph(text):
    t = re.sub(r'\s+', ' ', text).strip()
    if not t or "The following table:" in t:
        return True
    
    low_t = t.lower()
    
    if t.startswith('[') and t.endswith(']'):
        return True
        
    target_meta_words = ["collocation", "vocabulary reading", "vocabulary & reading", "vocabulary/reading", "chunking"]
    if any(word in low_t for word in target_meta_words):
        if len(t) < 60 or any(c in t for c in ['~', '[', ']', '▶', '■', '◆', '●', ':']):
            return True
            
    metadata_keywords = [
        "word arrangement", "fill in the blank", "교재 연계", "교재연계", 
        "객관-간접", "객관형", "간접형", "sentence transformation", 
        "correct sentence", "pg.", "page"
    ]
    if any(kw in low_t for kw in metadata_keywords):
        return True
        
    if any(t.startswith(char) for char in ['[', '▶', '■', '◆', '●', '★']):
        if any(kw in low_t for kw in ["vocabulary", "reading", "chapter", "level"]):
            return True
                
    return False

def apply_custom_style(run, font_name="Noto Sans KR", color_rgb=None):
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = color_rgb
        
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_paragraph_spacing(p, line=240, after=0, line_rule="auto"):
    """단락 줄간격/간격 설정"""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), line_rule)

def set_table_width_to_column(table):
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
        
    for child in list(tblPr):
        if child.tag.endswith('tblW'):
            tblPr.remove(child)
            
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def set_cell_properties(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcW'):
            tcPr.remove(child)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '5000')
    tcW.set(qn('w:type'), 'pct')
    tcPr.append(tcW)
    
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', 140), ('bottom', 140), ('left', 160), ('right', 160)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def apply_exam_page_layout(target_doc):
    """
    일반용 문서(2단 컬럼, 좁은 여백) 페이지 레이아웃 적용.
    pgSz: A4, pgMar: 상하좌우 720 (0.5인치), cols: 2단
    """
    body = target_doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)

    # 기존 pgSz, pgMar, cols 제거 후 재설정
    for tag in ['w:pgSz', 'w:pgMar', 'w:cols', 'w:docGrid']:
        el = sectPr.find(qn(tag))
        if el is not None:
            sectPr.remove(el)

    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    pgSz.set(qn('w:orient'), 'portrait')
    sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '720')
    pgMar.set(qn('w:right'), '720')
    pgMar.set(qn('w:bottom'), '720')
    pgMar.set(qn('w:left'), '720')
    pgMar.set(qn('w:header'), '567')
    pgMar.set(qn('w:footer'), '454')
    sectPr.append(pgMar)

    cols = OxmlElement('w:cols')
    cols.set(qn('w:equalWidth'), '0')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:sep'), '1')
    col1 = OxmlElement('w:col')
    col1.set(qn('w:space'), '425')
    col1.set(qn('w:w'), '5020')
    cols.append(col1)
    col2 = OxmlElement('w:col')
    col2.set(qn('w:space'), '0')
    col2.set(qn('w:w'), '5020')
    cols.append(col2)
    sectPr.append(cols)

    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:linePitch'), '360')
    sectPr.append(docGrid)

def flush_passage_buffer(target_doc, buffer):
    """버퍼링된 평문 지문들을 일반용 지문 상자(표 1칸) 구조로 빌드하여 삽입합니다."""
    if not buffer:
        return
        
    table = target_doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    set_table_width_to_column(table)
    cell = table.cell(0, 0)
    set_cell_properties(cell)
    
    cell.text = ""
    is_first = True
    
    for p_src in buffer:
        if is_first:
            p_dst = cell.paragraphs[0]
            is_first = False
        else:
            p_dst = cell.add_paragraph()
            
        p_dst.paragraph_format.line_spacing = 1.2
        p_dst.paragraph_format.space_after = Pt(4)
        
        if p_src.runs:
            for run in p_src.runs:
                dst_run = p_dst.add_run(run.text)
                dst_run.bold = run.bold
                dst_run.italic = run.italic
                dst_run.underline = run.underline
                if "정답" in run.text or is_originally_red(run):
                    apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=RGBColor(255, 0, 0))
                else:
                    apply_custom_style(dst_run, font_name="Noto Sans KR")
        else:
            dst_run = p_dst.add_run(p_src.text)
            if "정답" in p_src.text:
                apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=RGBColor(255, 0, 0))
            else:
                apply_custom_style(dst_run, font_name="Noto Sans KR")
                
    buffer.clear()


# ==========================================
# [서식 상수] 예시 문서 기준
# ==========================================
FONT_SIZE_QUESTION = Pt(11)    # 139700 EMU = 11pt  → 문제번호, 문제 지문 행
FONT_SIZE_CHOICE   = Pt(9)     # 114300 EMU = 9pt   → 선지, 정답, 일반 텍스트
LINE_SPACING_Q     = 276       # 문제 번호/지문 줄간격 (twips)
LINE_SPACING_BODY  = 240       # 선지/정답 줄간격


# ==========================================
# 1. 일반용 ➡️ 시험지용 변환 엔진
# ==========================================
def convert_general_to_exam_integrated(source_doc):
    template_path = "template_exam.docx"
    target_doc = Document(template_path) if os.path.exists(template_path) else Document()
    
    q_counter = 1
    red_color = RGBColor(255, 0, 0)
    
    for element in source_doc.element.body:
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, source_doc)
            p_text = p.text.strip()
            
            if should_skip_paragraph(p_text):
                continue
            
            is_p_answer = "정답" in p_text
            
            # 문제 번호로 시작하는 문단
            if re.match(r'^\d+[\.\s]', p_text):
                if q_counter > 1:
                    target_doc.add_paragraph()
                    target_doc.add_paragraph()
                
                new_p = target_doc.add_paragraph()
                # 문제번호 줄간격
                set_paragraph_spacing(new_p, line=LINE_SPACING_Q, after=0)
                
                clean_p_text = re.sub(r'^\d+\.\s*', '', p_text)
                clean_p_text = re.sub(r'^\d+\s+', '', clean_p_text)
                
                # 문제 번호 run: 볼드, 11pt
                run_num = new_p.add_run(f"{q_counter}.")
                run_num.bold = True
                run_num.font.size = FONT_SIZE_QUESTION
                apply_custom_style(run_num, font_name="Noto Sans KR",
                                   color_rgb=red_color if is_p_answer else None)
                
                # 번호 뒤 공백 run: 볼드, 9pt
                run_sp = new_p.add_run(" ")
                run_sp.bold = True
                run_sp.font.size = FONT_SIZE_CHOICE
                apply_custom_style(run_sp, font_name="Noto Sans KR")
                
                q_counter += 1
                
                if p.runs:
                    for run in p.runs:
                        r_text = run.text
                        if r_text.strip().startswith(p_text[:2]):
                            r_text = re.sub(r'^\d+\.\s*', '', r_text)
                            r_text = re.sub(r'^\d+\s+', '', r_text)
                        if r_text:
                            new_run = new_p.add_run(r_text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = FONT_SIZE_CHOICE
                            if is_p_answer or "정답" in run.text or is_originally_red(run):
                                apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=red_color)
                            else:
                                apply_custom_style(new_run, font_name="Noto Sans KR")
                else:
                    new_run = new_p.add_run(clean_p_text)
                    new_run.font.size = FONT_SIZE_CHOICE
                    apply_custom_style(new_run, font_name="Noto Sans KR",
                                       color_rgb=red_color if is_p_answer else None)
            else:
                # 선지, 정답, 지문 등 일반 단락
                new_p = target_doc.add_paragraph()
                set_paragraph_spacing(new_p, line=LINE_SPACING_BODY, after=0)
                
                if p.runs:
                    for run in p.runs:
                        new_run = new_p.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = FONT_SIZE_CHOICE
                        if is_p_answer or "정답" in run.text or is_originally_red(run):
                            apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=red_color)
                        else:
                            apply_custom_style(new_run, font_name="Noto Sans KR")
                else:
                    new_run = new_p.add_run(p.text)
                    new_run.font.size = FONT_SIZE_CHOICE
                    apply_custom_style(new_run, font_name="Noto Sans KR",
                                       color_rgb=red_color if is_p_answer else None)
                        
        elif element.tag.endswith('tbl'):
            src_tbl = docx.table.Table(element, source_doc)
            has_valid_content = False
            for row in src_tbl.rows:
                for cell in row.cells:
                    for p_cell in cell.paragraphs:
                        if p_cell.text.strip() and not should_skip_paragraph(p_cell.text):
                            has_valid_content = True
            if not has_valid_content:
                continue
            
            dst_tbl = target_doc.add_table(rows=len(src_tbl.rows), cols=len(src_tbl.columns))
            dst_tbl.style = 'Table Grid'
            set_table_width_to_column(dst_tbl)
            
            for r_idx, row in enumerate(src_tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    dst_cell = dst_tbl.cell(r_idx, c_idx)
                    set_cell_properties(dst_cell)
                    dst_cell.text = "" 
                    
                    is_first_p = True
                    for p_cell in cell.paragraphs:
                        if should_skip_paragraph(p_cell.text):
                            continue
                        if is_first_p:
                            dst_p_cell = dst_cell.paragraphs[0]
                            is_first_p = False
                        else:
                            dst_p_cell = dst_cell.add_paragraph()
                            
                        set_paragraph_spacing(dst_p_cell, line=LINE_SPACING_BODY, after=0)
                        
                        is_cell_answer = "정답" in p_cell.text
                        cell_color = red_color if is_cell_answer else None
                        
                        if p_cell.runs:
                            for run in p_cell.runs:
                                dst_run = dst_p_cell.add_run(run.text)
                                dst_run.bold = run.bold
                                dst_run.italic = run.italic
                                dst_run.underline = run.underline
                                dst_run.font.size = FONT_SIZE_CHOICE
                                if is_cell_answer or "정답" in run.text or is_originally_red(run):
                                    apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=red_color)
                                else:
                                    apply_custom_style(dst_run, font_name="Noto Sans KR")
                        else:
                            dst_run = dst_p_cell.add_run(p_cell.text)
                            dst_run.font.size = FONT_SIZE_CHOICE
                            apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=cell_color)

    # ★ 페이지 레이아웃 적용 (2단 컬럼, 좁은 여백)
    apply_exam_page_layout(target_doc)

    b_io = io.BytesIO()
    target_doc.save(b_io)
    return b_io.getvalue()


# ==========================================
# 2. 시험지용 ➡️ 일반용 변환 엔진
# ==========================================
def convert_exam_to_general_integrated(source_doc):
    target_doc = Document()
    
    q_counter = 1
    red_color = RGBColor(255, 0, 0)
    
    inside_question = False
    passage_buffer = []
    
    for element in source_doc.element.body:
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, source_doc)
            p_text = p.text.strip()
            
            if should_skip_paragraph(p_text):
                continue
                
            if re.match(r'^\d+[\.\s]', p_text):
                flush_passage_buffer(target_doc, passage_buffer)
                inside_question = True
                
                new_p = target_doc.add_paragraph()
                new_p.paragraph_format.space_before = p.paragraph_format.space_before
                new_p.paragraph_format.space_after = p.paragraph_format.space_after
                
                clean_p_text = re.sub(r'^\d+\.\s*', '', p_text)
                clean_p_text = re.sub(r'^\d+\s+', '', clean_p_text)
                
                run_num = new_p.add_run(f"{q_counter}.")
                run_num.bold = True
                apply_custom_style(run_num, font_name="Noto Sans KR", color_rgb=red_color if "정답" in p_text else None)
                q_counter += 1
                
                # 공백 run
                run_sp = new_p.add_run(" ")
                run_sp.bold = True
                apply_custom_style(run_sp, font_name="Noto Sans KR")
                
                if p.runs:
                    for run in p.runs:
                        r_text = run.text
                        if r_text.strip().startswith(p_text[:2]):
                            r_text = re.sub(r'^\d+\.\s*', '', r_text)
                            r_text = re.sub(r'^\d+\s+', '', r_text)
                            
                        if r_text:
                            new_run = new_p.add_run(r_text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if "정답" in run.text or is_originally_red(run):
                                apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=red_color)
                            else:
                                apply_custom_style(new_run, font_name="Noto Sans KR")
                else:
                    new_run = new_p.add_run(clean_p_text)
                    apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=red_color if "정답" in p_text else None)
            
            else:
                is_option_or_answer = any(idx in p_text for idx in ['①', '②', '③', '④', '⑤']) or "정답" in p_text
                
                if inside_question and not is_option_or_answer:
                    passage_buffer.append(p)
                else:
                    if is_option_or_answer:
                        flush_passage_buffer(target_doc, passage_buffer)
                    
                    new_p = target_doc.add_paragraph()
                    new_p.paragraph_format.space_before = p.paragraph_format.space_before
                    new_p.paragraph_format.space_after = p.paragraph_format.space_after
                    
                    is_p_answer = "정답" in p_text
                    current_color = red_color if is_p_answer else None
                    
                    if p.runs:
                        for run in p.runs:
                            new_run = new_p.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if is_p_answer or "정답" in run.text or is_originally_red(run):
                                apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=red_color)
                            else:
                                apply_custom_style(new_run, font_name="Noto Sans KR")
                    else:
                        new_run = new_p.add_run(p.text)
                        apply_custom_style(new_run, font_name="Noto Sans KR", color_rgb=current_color)
                        
        elif element.tag.endswith('tbl'):
            # ★ 시험지용 → 일반용: 표(지문 상자)를 passage_buffer로 처리
            src_tbl = docx.table.Table(element, source_doc)
            
            # 표 내 유효한 내용이 있는지 확인
            has_valid_content = any(
                p_cell.text.strip() and not should_skip_paragraph(p_cell.text)
                for row in src_tbl.rows
                for cell in row.cells
                for p_cell in cell.paragraphs
            )
            if not has_valid_content:
                continue

            if inside_question:
                # 지문 상자: 표 내 단락들을 passage_buffer에 추가
                for row in src_tbl.rows:
                    for cell in row.cells:
                        for p_cell in cell.paragraphs:
                            if p_cell.text.strip() and not should_skip_paragraph(p_cell.text):
                                passage_buffer.append(p_cell)
            else:
                # 문제 바깥의 표는 그대로 복사
                flush_passage_buffer(target_doc, passage_buffer)
                dst_tbl = target_doc.add_table(rows=len(src_tbl.rows), cols=len(src_tbl.columns))
                dst_tbl.style = 'Table Grid'
                set_table_width_to_column(dst_tbl)
                
                for r_idx, row in enumerate(src_tbl.rows):
                    for c_idx, cell in enumerate(row.cells):
                        dst_cell = dst_tbl.cell(r_idx, c_idx)
                        set_cell_properties(dst_cell)
                        dst_cell.text = ""
                        
                        is_first_p = True
                        for p_cell in cell.paragraphs:
                            if should_skip_paragraph(p_cell.text):
                                continue
                            if is_first_p:
                                dst_p_cell = dst_cell.paragraphs[0]
                                is_first_p = False
                            else:
                                dst_p_cell = dst_cell.add_paragraph()
                                
                            dst_p_cell.paragraph_format.line_spacing = 1.2
                            dst_p_cell.paragraph_format.space_after = Pt(2)
                            
                            is_cell_answer = "정답" in p_cell.text
                            cell_color = red_color if is_cell_answer else None
                            
                            if p_cell.runs:
                                for run in p_cell.runs:
                                    dst_run = dst_p_cell.add_run(run.text)
                                    dst_run.bold = run.bold
                                    dst_run.italic = run.italic
                                    dst_run.underline = run.underline
                                    if is_cell_answer or "정답" in run.text or is_originally_red(run):
                                        apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=red_color)
                                    else:
                                        apply_custom_style(dst_run, font_name="Noto Sans KR")
                            else:
                                dst_run = dst_p_cell.add_run(p_cell.text)
                                apply_custom_style(dst_run, font_name="Noto Sans KR", color_rgb=cell_color)

    flush_passage_buffer(target_doc, passage_buffer)

    b_io = io.BytesIO()
    target_doc.save(b_io)
    return b_io.getvalue()


# ==========================================
# 3. Streamlit 웹 대시보드 인터페이스
# ==========================================
st.title("📝 정기평가 양식 최고 고도화 시스템")
st.markdown("가로 너비 **100% 자동 동기화** / **Noto Sans KR 글꼴** / **정답 추적 빨간색 고정** / **양방향 지문 완벽 보존** 버전입니다.")

uploaded_file = st.file_uploader("변환할 정기평가 Word 파일(.docx)을 업로드하세요.", type=["docx"])

if uploaded_file is not None:
    doc = Document(uploaded_file)
    
    conversion_mode = st.radio(
        "원하시는 변환 작업 방향을 선택하세요:",
        [
            "일반용 ➡️ 시험지용 (2단 레이아웃 단 맞춤, 태그 소거, 문제 간 2줄 줄바꿈 적용)", 
            "시험지용 ➡️ 일반용 (1단 레이아웃 복원, 평문 영어 지문을 상자 표 내부로 자동 복구 추출)"
        ]
    )
    
    if st.button("🚀 선택한 모드로 정밀 변환 시작", use_container_width=True):
        with st.spinner("문서 컨텍스트 및 지문 인코딩 예외 교정 중..."):
            try:
                if "일반용 ➡️ 시험지용" in conversion_mode:
                    out_bytes = convert_general_to_exam_integrated(doc)
                    download_filename = "정기평가_최종형_시험지문서.docx"
                else:
                    out_bytes = convert_exam_to_general_integrated(doc)
                    download_filename = "정기평가_복원형_일반문서.docx"
                    
                st.success("🎉 지문 자동 상자 래핑 및 문제 간격 조정 패치가 적용되었습니다!")
                st.download_button(
                    label="💾 완성본 파일 다운로드",
                    data=out_bytes,
                    file_name=download_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"⚠️ 시스템 변환 처리 중 오류 발생: {str(e)}")
